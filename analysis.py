import pandas as pd
import matplotlib.pyplot as plt
from pathlib import Path
import seaborn as sns
from datetime import datetime
import logging
from docx import Document
import re
import numpy as np
import traceback

from util import clean_title

plt.rcParams['font.sans-serif'] = ['SimHei']  # 设置中文字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题


def analyze_post_trends(csv_path: str = "data.csv") -> None:
    """分析文章发表趋势并生成报表"""
    logging.info("开始分析文章发表趋势")
    
    # 读取CSV文件
    df = pd.read_csv(csv_path)
    
    # 将发表时间转换为datetime类型
    df['发表时间'] = pd.to_datetime(df['发表时间'])
    
    # 添加年月列
    df['年月'] = df['发表时间'].dt.strftime('%Y-%m')
    
    # 按年月和板块统计文章数量
    monthly_counts = df.groupby(['年月', '板块']).size().unstack(fill_value=0)
    total_monthly = df.groupby('年月').size()
    
    # 创建输出目录
    output_dir = Path("./分析报告")
    output_dir.mkdir(exist_ok=True)
    
    # 生成总体趋势图
    plt.figure(figsize=(15, 8))
    total_monthly.plot(kind='line', marker='o')
    plt.title('文章发表数量月度趋势')
    plt.xlabel('年月')
    plt.ylabel('文章数量')
    plt.xticks(rotation=45)
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(output_dir / '总体发表趋势.png')
    plt.close()
    
    # 生成堆叠面积图
    plt.figure(figsize=(15, 8))
    monthly_counts.plot(kind='area', stacked=True)
    plt.title('各板块文章发表数量趋势')
    plt.xlabel('年月')
    plt.ylabel('文章数量')
    plt.xticks(rotation=45)
    plt.legend(title='板块', bbox_to_anchor=(1.05, 1), loc='upper left')
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(output_dir / '板块分布趋势.png')
    plt.close()
    
    # 生成热力图
    plt.figure(figsize=(15, 8))
    sns.heatmap(monthly_counts.T, cmap='YlOrRd', annot=True, fmt='g')
    plt.title('文章发表数量热力图')
    plt.xlabel('年月')
    plt.ylabel('板块')
    plt.tight_layout()
    plt.savefig(output_dir / '发表数量热力图.png')
    plt.close()
    
    # 生成统计报告
    with open(output_dir / '统计报告.txt', 'w', encoding='utf-8') as f:
        f.write('文章发表统计报告\n')
        f.write('=' * 50 + '\n\n')
        
        f.write('1. 总体统计\n')
        f.write(f'总文章数：{len(df)}\n')
        f.write(f'统计周期：{df["发表时间"].min().strftime("%Y-%m")} 至 {df["发表时间"].max().strftime("%Y-%m")}\n\n')
        
        f.write('2. 板块分布\n')
        block_stats = df['板块'].value_counts()
        for block, count in block_stats.items():
            f.write(f'{block}: {count}篇 ({count/len(df)*100:.2f}%)\n')
        f.write('\n')
        
        f.write('3. 月度发表TOP5\n')
        top_months = total_monthly.sort_values(ascending=False).head()
        for month, count in top_months.items():
            f.write(f'{month}: {count}篇\n')
            
    logging.info("分析完成，报告已生成")

def analyze_post_quality(csv_path: str = "data.csv", output_dir='分析报告') -> None:
    """分析近期文章并生成报表"""
    logging.info("开始分析近期文章")
    
    # 定义标准化函数
    def normalize(series):
        if series.max() == series.min():
            return series * 0  # 如果所有值相同，返回0
        return (series - series.min()) / (series.max() - series.min())
    
    # 将 output_dir 转换为 Path 对象
    output_dir = Path(output_dir)
    
    # 确保输出目录存在
    output_dir.mkdir(exist_ok=True)
    
    # 读取CSV文件
    df = pd.read_csv(csv_path)
    
    # 确保数值列为数值类型
    df['浏览数'] = pd.to_numeric(df['浏览数'], errors='coerce')
    df['点赞数'] = pd.to_numeric(df['点赞数'], errors='coerce')
    df['收藏数'] = pd.to_numeric(df['收藏数'], errors='coerce')
    df['字数'] = pd.to_numeric(df['字数'], errors='coerce')
    
    # 将NaN值替换为0
    df[['浏览数', '点赞数', '收藏数']] = df[['浏览数', '点赞数', '收藏数']].fillna(0)
    
    # 清理标题
    df['标题'] = df['标题'].apply(clean_title)
    
    # 过滤数据
    df = df[
        (df['作者'] != 'Admin_荆棘鸟') &  # 过滤管理员文章
        (df['字数'] > 0)  # 过滤字数为0的文章
    ]
    
    # 数据预处理
    df['浏览数'] = df['浏览数'].astype(int)
    df['点赞数'] = df['点赞数'].astype(int)
    df['收藏数'] = df['收藏数'].astype(int)
    
    # 设置字数阈值和权重计算参数
    MIN_WORDS = 5000      # 最低字数要求
    BASE_WORDS = 15000    # 基准字数
    TARGET_WORDS = 30000  # 目标字数
    MAX_WEIGHT = 2.0      # 最高字数权重
    
    # 计算字数权重（强化高字数的权重）
    def calculate_word_weight(words):
        if words < MIN_WORDS:
            return 0.8  # 最低保底权重
        elif words < BASE_WORDS:
            # 5000-15000字区间
            ratio = (words - MIN_WORDS) / (BASE_WORDS - MIN_WORDS)
            return 0.8 + (0.4 * ratio)  # 0.8 到 1.2 的过渡
        elif words < TARGET_WORDS:
            # 15000-30000字区间
            ratio = (words - BASE_WORDS) / (TARGET_WORDS - BASE_WORDS)
            return 1.2 + (0.4 * ratio)  # 1.2 到 1.6 的过渡
        else:
            # 30000字以上，继续线性增长但设置上限
            extra_words = words - TARGET_WORDS
            extra_weight = min(0.4, extra_words / 10000 * 0.2)  # 每增加10000字增加0.2的权重，最多增加0.4
            return 1.6 + extra_weight  # 最高不超过2.0
    
    # 应用字数权重
    df['字数权重'] = df['字数'].apply(calculate_word_weight)
    
    # 1. 计算基础互动率
    df['互动率'] = (df['点赞数'] + df['收藏数']) / df['浏览数']
    
    # 2. 修改字数权重计算方式，对短文章施加惩罚
    df['字数权重'] = df['字数'].apply(lambda x: min(1.0, (x / MIN_WORDS) ** 0.5))
    
    # 3. 标准化各个指标
    df['浏览数_标准化'] = normalize(df['浏览数'])
    df['互动率_标准化'] = normalize(df['互动率'])
    
    # 4. 标准化函数
    def normalize(series):
        if series.max() == series.min():
            return series * 0  # 如果所有值相同，返回0
        return (series - series.min()) / (series.max() - series.min())
    
    # 5. 标准化各个指标
    df['浏览数_标准化'] = normalize(df['浏览数'])
    df['互动率_标准化'] = normalize(df['互动率'])
    
    # 修改评分算法部分
    # 1. 对浏览数进行对数转换，减少极值影响
    df['浏览数_对数'] = np.log1p(df['浏览数'])  # log1p 避免 log(0)
    
    # 2. 重新定义互动率计算
    df['互动率'] = (df['点赞数'] + df['收藏数']) / df['浏览数'].clip(lower=1)
    
    # 3. 设置最小阈值，防止过度奖励低浏览量文章
    MIN_VIEWS = 100  # 最小浏览量阈值
    df['浏览量权重'] = df['浏览数'].apply(lambda x: min(1.0, (x / MIN_VIEWS) ** 0.5))
    
    # 4. 标准化各指标
    df['浏览数_标准化'] = normalize(df['浏览数_对数'])
    df['互动率_标准化'] = normalize(df['互动率'])
    
    # 3. 重新计算互动质量分
    df['互动质量'] = (df['收藏数'] * 2 + df['点赞数'] + df['评论数']) / df['浏览数'].clip(lower=1)
    df['互动质量_标准化'] = normalize(df['互动质量'])
    
    # 2. 计算互动密度（考虑字数的互动频率）
    WORD_SEGMENT = 1000  # 每1000字为一个段
    df['段落数'] = (df['字数'] / WORD_SEGMENT).clip(lower=1)
    df['互动密度'] = (df['点赞数'] + df['收藏数']) / df['段落数']
    
    # 确保互动密度列存在
    if '互动密度' not in df.columns:
        df['互动密度'] = 0  # 如果没有计算出互动密度，初始化为0

    # 标准化互动密度
    df['互动密度_标准化'] = normalize(df['互动密度'])
    
    # 1. 增加评论数的处理
    df['评论数'] = pd.to_numeric(df['评论数'], errors='coerce').fillna(0)
    
    # 2. 增加时间因素的考虑
    df['发表时间'] = pd.to_datetime(df['发表时间'])
    # 使用指数衰减函数计算时间权重
    df['时间衰减因子'] = np.exp(-0.01 * (datetime.now() - df['发表时间']).dt.days)
    
    # 4. 重新计算综合评分，加入时间衰减因子
    df['综合评分'] = (
        df['浏览数_标准化'] * 0.30 +     # 增加浏览量权重
        df['互动质量_标准化'] * 0.15 +   # 保持互动质量权重
        df['互动密度_标准化'] * 0.10 +   # 增加互动密度权重
        df['时间衰减因子'] * 0.2 +      # 保持时间衰减权重
        df['字数权重'] * 0.25           # 保持字数权重
    )
    
    # 使用字数权重的二次方作为最终调节因子，进一步强化字数的影响
    df['综合评分'] = df['综合评分'] * (df['字数权重'] ** 2)

    # 4. 在输出中添加新的指标
    all_articles = df[['标题', '板块', '字数', '浏览数', '点赞数', '收藏数', 
                      '互动率', '互动质量', '互动密度',
                      '综合评分']].sort_values('综合评分', ascending=False)
    
    # 添加排名列
    all_articles.insert(0, '排名', range(1, len(all_articles) + 1))
    
    # 格式化数值列
    all_articles['互动率'] = all_articles['互动率'].map('{:.2%}'.format)
    all_articles['互动质量'] = all_articles['互动质量'].map('{:.2e}'.format)
    all_articles['互动密度'] = all_articles['互动密度'].map('{:.2f}'.format)
    all_articles['综合评分'] = all_articles['综合评分'].map('{:.4f}'.format)

    all_articles.to_csv(output_dir / '近期文章推荐排名.csv', index=False, encoding='utf-8-sig')
    
    # 生成板块质量分析
    block_quality = df.groupby('板块').agg({
        '浏览数': ['mean', 'median', 'std', 'count'],  # 添加更多统计指标
        '点赞数': ['mean', 'median', 'std'],
        '收藏数': ['mean', 'median', 'std'],
        '互动率': ['mean', 'median', 'std'],
        '字数': ['mean', 'median', 'std'],
        '综合评分': ['mean', 'median', 'std']
    }).round(2)
    
    # 重命名列名使其更易读
    block_quality.columns = [
        f'{col[0]}_{col[1]}'.replace('_mean', '_平均值')
                           .replace('_median', '_中位数')
                           .replace('_std', '_标准差')
                           .replace('_count', '_总数') 
        for col in block_quality.columns
    ]
    
    # 添加百分比统计
    total_posts = len(df)
    block_quality['占比'] = (block_quality['浏览数_总数'] / total_posts * 100).round(2).astype(str) + '%'
    
    block_quality.to_csv(output_dir / '板块质量分析.csv', encoding='utf-8-sig')
    
    # 生成可视化图表
    plt.figure(figsize=(15, 8))
    
    # 1. 互动率与浏览数的散点图
    plt.scatter(df['浏览数'], df['互动率'], alpha=0.5)
    plt.title('文章浏览数与互动率关系')
    plt.xlabel('浏览数')
    plt.ylabel('互动率')
    
    # 添加TOP10文章标注
    top_10 = df.nlargest(10, '综合评分')
    for _, row in top_10.iterrows():
        plt.annotate(row['标题'][:10] + '...', 
                    (row['浏览数'], row['互动率']),
                    xytext=(5, 5), textcoords='offset points')
    
    plt.tight_layout()
    plt.savefig(output_dir / '浏览数与互动率关系.png')
    plt.close()
    
    # 在计算完文章评分后，添加作者分析部分
    
    # 1. 计算作者级别的统计数据
    author_stats = df.groupby('作者').agg({
        '综合评分': ['mean', 'std', 'count'],
        '字数': 'sum',
        '浏览数': 'sum',
        '点赞数': 'sum',
        '收藏数': 'sum'
    })
    
    # 重命名列
    author_stats.columns = ['平均文章评分', '评分标准差', '文章数量', 
                          '总字数', '总浏览数', '总点赞数', '总收藏数']
    
    # 2. 计算作者评分指标
    author_stats['内容产出力'] = normalize(author_stats['总字数']) * 0.7 + \
                             normalize(author_stats['文章数量']) * 0.3
    
    author_stats['互动影响力'] = normalize(author_stats['总浏览数']) * 0.4 + \
                             normalize(author_stats['总点赞数']) * 0.3 + \
                             normalize(author_stats['总收藏数']) * 0.3
    
    # 计质量稳定性（满分1分，标准差越大扣分越多）
    author_stats['质量稳定性'] = 1 - normalize(author_stats['评分标准差'])
    
    # 3. 计算综合作者评分
    author_stats['作者评分'] = (
        author_stats['平均文章评分'] * 0.4 +    # 文章质量权重
        author_stats['内容产出力'] * 0.25 +     # 产出力权重
        author_stats['互动影响力'] * 0.25 +     # 影响力权重
        author_stats['质量稳定性'] * 0.1        # 稳定性权重
    )
    
    # 将作者评分标准化到百分制
    author_stats['作者评分'] = (author_stats['作者评分'] / author_stats['作者评分'].max()) * 100
    
    # 4. 生成作者推荐排名
    author_ranking = author_stats.sort_values('作者评分', ascending=False)
    
    # 只保留发文量达到要求的作者
    MIN_ARTICLES = 3  # 最少文章数要求
    author_ranking = author_ranking[author_ranking['文章数量'] >= MIN_ARTICLES]
    
    # 添加排名列
    author_ranking.insert(0, '排名', range(1, len(author_ranking) + 1))
    
    # 格式化输出数据
    output_columns = ['排名', '文章数量', '平均文章评分', '总字数', 
                     '总浏览数', '总点赞数', '总收藏数', '质量稳定性', '作者评分']
    
    author_ranking = author_ranking[output_columns].round(2)
    
    # 保存作者排名
    author_ranking.to_csv(output_dir / '作者推荐排名.csv', encoding='utf-8-sig')
    
    # 生成作者质量分布图
    plt.figure(figsize=(12, 6))
    plt.scatter(author_ranking['文章数量'], author_ranking['平均文章评分'], 
               alpha=0.5, s=author_ranking['总浏览数']/1000)
    
    # 标注TOP5作者
    top_5_authors = author_ranking.head()
    for idx, row in top_5_authors.iterrows():
        plt.annotate(idx, 
                    (row['文章数量'], row['平均文章评分']),
                    xytext=(5, 5), textcoords='offset points')
    
    plt.title('作者文章数量与质量分布')
    plt.xlabel('文章数量')
    plt.ylabel('平均文章评分')
    plt.tight_layout()
    plt.savefig(output_dir / '作者质量分布.png')
    plt.close()
    
    # 5. 为每个作者找出代表作
    def get_representative_works(author_articles, top_n=3):
        """获取作者的代表作"""
        # 按综合评分排序选择前N篇
        top_articles = author_articles.nlargest(top_n, '综合评分')
        
        # 格式化输出信息
        return [f"{row['标题']} (评分:{float(row['综合评分']):.2f}, 浏览:{int(row['浏览数'])}, 点赞:{int(row['点赞数'])})"
                for _, row in top_articles.iterrows()]
    
    # 创建作者代表作字典
    author_top_works = {}
    for author in author_ranking.index:
        author_articles = df[df['作者'] == author]
        author_top_works[author] = get_representative_works(author_articles)
    
    # 创建详细的作者分析报告
    doc = Document()
    doc.add_heading('作者分析报告', 0)
    
    # 添加总体统计信息
    doc.add_heading('1. 总体统计', level=1)
    doc.add_paragraph(f'分析周期：{df["发表时间"].min().strftime("%Y-%m-%d")} 至 {df["发表时间"].max().strftime("%Y-%m-%d")}')
    doc.add_paragraph(f'作者总数：{len(author_ranking)}')
    doc.add_paragraph(f'文章总数：{len(df)}')
    
    # 添加TOP作者详细分析
    doc.add_heading('2. TOP10作者详细分析', level=1)
    
    for author in author_ranking.head(10).index:
        doc.add_heading(f'作者：{author}', level=2)
        stats = author_ranking.loc[author]
        
        # 基础统计信息
        doc.add_paragraph(f'排名：第{int(stats["排名"])}名')
        doc.add_paragraph(f'文章数量：{int(stats["文章数量"])}篇')
        doc.add_paragraph(f'总字数：{int(stats["总字数"]):,}字')
        doc.add_paragraph(f'平均文章评分：{float(stats["平均文章评分"]):.2f}')
        doc.add_paragraph(f'质量稳定性：{float(stats["质量稳定性"]):.2f}')
        
        # 互动数据
        doc.add_paragraph('互动数据：')
        doc.add_paragraph(f'- 总浏览数：{int(stats["总浏览数"]):,}')
        doc.add_paragraph(f'- 总点赞数：{int(stats["总点赞数"]):,}')
        doc.add_paragraph(f'- 总收藏数：{int(stats["总收藏数"]):,}')
        
        # 代表作列表
        doc.add_paragraph('代表作品：')
        for i, work in enumerate(author_top_works[author], 1):
            doc.add_paragraph(f'{i}. {work}', style='List Number')
        
        # 添加分隔线
        doc.add_paragraph('=' * 50)
    
    # 保存详细报告
    doc.save(output_dir / '作者分析详细报告.docx')
    
    # 将代表作信息添加到作者排名CSV中
    author_ranking['代表作品'] = author_ranking.index.map(
        lambda x: ' || '.join(author_top_works.get(x, ['无']))
    )
    
    # 更新输出列
    output_columns = ['排名', '文章数量', '平均文章评分', '总字数', 
                     '总浏览数', '总点赞数', '总收藏数', 
                     '质量稳定性', '作者评分', '代表作品']
    
    # 保存作者排名
    author_ranking[output_columns].to_csv(
        output_dir / '作者推荐排名.csv', 
        encoding='utf-8-sig'
    )
    
    # 输出统计结果
    logging.info(f"分析完成，结果已保存到 {output_dir} 目录")
    

if __name__ == "__main__":
    # 设置日志配置
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    # 设置数据文件路径
    data_file = "data.csv"
    try:
        # 依次执行所有分析函数
        analyze_post_trends(data_file)
        analyze_post_quality(data_file)
        logging.info("所有分析任务已完成")
    except Exception as e:
        logging.error(f"分析过程中出现错误: {str(e)}")
        logging.error(f"错误堆栈:\n{traceback.format_exc()}")
            