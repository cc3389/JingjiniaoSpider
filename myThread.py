import os.path
import re
from bs4 import BeautifulSoup
from decode import decode_base64_in_js
from util import *
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn
# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)




def extract_title_and_counts(soup):
    """
    提取标题和推荐、收藏数
    """
    title = soup.select_one('#thread_subject').text
    recommend_num = soup.find(id='recommendv_add').text
    favorite_num = soup.find(id='favoritenumber').text
    return title, recommend_num, favorite_num




def get_last_page_num(soup):
    """
    获取最后一页的页码
    """
    last_page_tag = soup.find("span", title=re.compile("共 [0-9]+ 页"))
    return int(last_page_tag.text.replace(" ", "").replace("/", "").replace("页", "")) if last_page_tag else 1


def process_tags(t_f, document, download_images):
    """
    处理标签并按顺序添加文本和图片到文档
    """
    word_count = 0  # 添加字数计数器
    for tags in t_f:
        temp_text = ''
        for tag in tags:
            if tag.name == 'script':
                continue

            # 清理文本中的多余换行符和空白字符
            current_text = tag.text.strip()
            current_text = re.sub(r'\s+', ' ', current_text)  # 将多个空白字符替换为单个空格

            # 添加处理后的文本
            temp_text += current_text

            # 如果累积了文本，先添加到文档中
            if temp_text.strip():
                # 移除不合法的XML字符
                clean_text = ''.join(char for char in temp_text.strip() if ord(char) >= 32 or char in '\n\r\t')
                document.add_paragraph(clean_text)
                word_count += len(clean_text)  # 统计字数
                temp_text = ''

            # 处理图片
            if download_images:
                if hasattr(tag, 'find_all'):
                    for img in tag.find_all('img'):
                        try:
                            img_url = 'https://www.jingjiniao.info/' + img['file']
                            image_stream = download_image(img_url)
                            if image_stream:
                                pic = document.add_picture(image_stream)
                                # 获取图片原始宽高比
                                aspect_ratio = pic.height / pic.width
                                # 设置最大宽度为页面宽度的80%
                                max_width = Inches(6)  # A4纸宽度约为8.27英寸
                                if pic.width > max_width:
                                    pic.width = max_width
                                    pic.height = int(pic.width * aspect_ratio)
                        except:
                            pass

    return word_count  # 返回该部分的字数统计


def thread_spider(thread_url, block_name, download_images):
    """
    爬取具体的文章并存入文档中
    """
    start_time = time.time()
    document = Document()
    
    # 设置默认字体为宋体
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    page_num = 1
    total_word_count = 0
    
    try:
        while True:
            response = make_request(thread_url)
            if response is None:
                raise Exception('请求超时')
            soup = BeautifulSoup(response.text, 'html.parser')

            if len(soup.select("#messagetext")) > 0:
                error_message = soup.select("#messagetext")[0].text + soup.select("#messagetext")[0].next_sibling.text
                raise Exception(error_message)

            if page_num == 1:
                title, recommend_num, favorite_num = extract_title_and_counts(soup)
                title = clean_title(title)

            t_f = soup.select(".t_f div")
            if len(soup.select(".t_f script")) == 2:
                base64_js_str = soup.select(".t_f script")[1].text
                wmsj_enmessage_str = decode_base64_in_js(base64_js_str)
                wmsj_enmessageTag = BeautifulSoup(wmsj_enmessage_str, 'html.parser')
                t_f.insert(0, wmsj_enmessageTag)
            soup.find('div', class_='typeoption')
            for tags in t_f:
                for tag in tags.find_all(style='display:none') + tags.find_all(class_='jammer') \
                           + tags.find_all(['br', 'a', 'i']) + tags.find_all('div', class_='quote'):
                    tag.decompose()
            # 下载封面图片
            if download_images:
                img_tags = soup.select(".typeoption img")
                for img in img_tags:
                    try:
                        img_url = 'https://www.jingjiniao.info/' + img['src']
                        image_stream = download_image(img_url)
                        if image_stream:
                            pic = document.add_picture(image_stream)
                            # 获取图片原始宽高比
                            aspect_ratio = pic.height / pic.width
                            # 设置最大宽度为页面宽度的80%
                            max_width = Inches(6)  # A4纸宽度约为8.27英寸
                            if pic.width > max_width:
                                pic.width = max_width
                                pic.height = int(pic.width * aspect_ratio)
                    except:
                        pass
            # 获取文章内容
            page_word_count = process_tags(t_f, document, download_images)  # 获取每页的字数
            total_word_count += page_word_count  # 累加总字数
            nextLinkTag = soup.select('.nxt')
            if nextLinkTag:
                thread_url = nextLinkTag[0].attrs['href']
                page_num += 1
            else:
                break

        # 在保存文档前创建目录
        save_dir = os.path.join("./小说输出/", block_name)
        os.makedirs(save_dir, exist_ok=True)

        document.save(os.path.join(save_dir, clean_title(title) + ".docx"))

        end_time = time.time()
        logging.info(f'文章《{title}》爬取完成，总字数：{total_word_count}，总耗时: {end_time - start_time:.2f}秒')

    except Exception as e:
        logging.error(f'爬取 {thread_url}失败。原因： {e}')
        return 0, 0, 0  # 添加total_word_count的返回值

    return recommend_num, favorite_num, total_word_count


if __name__ == '__main__':
    recommend_num, favorite_num, total_word_count = thread_spider('https://www.jingjiniao.info/forum.php?mod=viewthread&tid=54677&page=1&authorid=18199', 'test', False)
    print(f'点赞数: {recommend_num}, 收藏数: {favorite_num}, 总字数: {total_word_count}')
