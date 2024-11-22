import logging

from docx import Document
from matplotlib import pyplot as plt

from util import normalize


def analyze_author(df, output_dir):
    """分析作者并生成报告"""
    # 1. 计算作者级别的统计数据
    author_stats = df.groupby('作者').agg({
        '综合评分': ['mean', 'std', 'count'],
        '字数': 'sum',
        '浏览数': 'sum',
        '点赞数': 'sum',
        '收藏数': 'sum'
    })

    # 重命名列
    author_stats.columns = ['平均文章评分', '评分标准差', '文章数量',
                          '总字数', '总浏览数', '总点赞数', '总收藏数']

    # 2. 计算作者评分指标
    author_stats['内容产出力'] = normalize(author_stats['总字数']) * 0.7 + \
                             normalize(author_stats['文章数量']) * 0.3

    author_stats['互动影响力'] = normalize(author_stats['总浏览数']) * 0.4 + \
                             normalize(author_stats['总点赞数']) * 0.3 + \
                             normalize(author_stats['总收藏数']) * 0.3

    # 计质量稳定性（满分1分，标准差越大扣分越多）
    author_stats['质量稳定性'] = 1 - normalize(author_stats['评分标准差'])

    # 3. 计算综合作者评分
    author_stats['作者评分'] = (
        author_stats['平均文章评分'] * 0.4 +    # 文章质量权重
        author_stats['内容产出力'] * 0.25 +     # 产出力权重
        author_stats['互动影响力'] * 0.25 +     # 影响力权重
        author_stats['质量稳定性'] * 0.1        # 稳定性权重
    )

    # 将作者评分标准化到百分制
    author_stats['作者评分'] = (author_stats['作者评分'] / author_stats['作者评分'].max()) * 100

    # 4. 生成作者推荐排名
    author_ranking = author_stats.sort_values('作者评分', ascending=False)

    # 只保留发文量达到要求的作者
    MIN_ARTICLES = 3  # 最少文章数要求
    author_ranking = author_ranking[author_ranking['文章数量'] >= MIN_ARTICLES]

    # 添加排名列
    author_ranking.insert(0, '排名', range(1, len(author_ranking) + 1))

    # 格式化输出数据
    output_columns = ['排名', '文章数量', '平均文章评分', '总字数',
                     '总浏览数', '总点赞数', '总收藏数', '质量稳定性', '作者评分']

    author_ranking = author_ranking[output_columns].round(2)

    # 保存作者排名
    author_ranking.to_csv(output_dir / '作者推荐排名.csv', encoding='utf-8-sig')

    # 生成作者质量分布图
    plt.figure(figsize=(12, 6))
    plt.scatter(author_ranking['文章数量'], author_ranking['平均文章评分'],
               alpha=0.5, s=author_ranking['总浏览数']/1000)

    # 标注TOP5作者
    top_5_authors = author_ranking.head()
    for idx, row in top_5_authors.iterrows():
        plt.annotate(idx,
                    (row['文章数量'], row['平均文章评分']),
                    xytext=(5, 5), textcoords='offset points')

    plt.title('作者文章数量与质量分布')
    plt.xlabel('文章数量')
    plt.ylabel('平均文章评分')
    plt.tight_layout()
    plt.savefig(output_dir / '作者质量分布.png')
    plt.close()

    # 5. 为每个作者找出代表作
    def get_representative_works(author_articles, top_n=3):
        """获取作者的代表作"""
        # 按综合评分排序选择前N篇
        top_articles = author_articles.nlargest(top_n, '综合评分')

        # 格式化输出信息
        return [f"{row['标题']} (评分:{float(row['综合评分']):.2f}, 浏览:{int(row['浏览数'])}, 点赞:{int(row['点赞数'])})"
                for _, row in top_articles.iterrows()]

    # 创建作者代表作字典
    author_top_works = {}
    for author in author_ranking.index:
        author_articles = df[df['作者'] == author]
        author_top_works[author] = get_representative_works(author_articles)

    # 创建详细的作者分析报告
    doc = Document()
    doc.add_heading('作者分析报告', 0)

    # 添加总体统计信息
    doc.add_heading('1. 总体统计', level=1)
    doc.add_paragraph(f'分析周期：{df["发表时间"].min().strftime("%Y-%m-%d")} 至 {df["发表时间"].max().strftime("%Y-%m-%d")}')
    doc.add_paragraph(f'作者总数：{len(author_ranking)}')
    doc.add_paragraph(f'文章总数：{len(df)}')

    # 添加TOP作者详细分析
    doc.add_heading('2. TOP10作者详细分析', level=1)

    for author in author_ranking.head(10).index:
        doc.add_heading(f'作者：{author}', level=2)
        stats = author_ranking.loc[author]

        # 基础统计信息
        doc.add_paragraph(f'排名：第{int(stats["排名"])}名')
        doc.add_paragraph(f'文章数量：{int(stats["文章数量"])}篇')
        doc.add_paragraph(f'总字数：{int(stats["总字数"]):,}字')
        doc.add_paragraph(f'平均文章评分：{float(stats["平均文章评分"]):.2f}')
        doc.add_paragraph(f'质量稳定性：{float(stats["质量稳定性"]):.2f}')

        # 互动数据
        doc.add_paragraph('互动数据：')
        doc.add_paragraph(f'- 总浏览数：{int(stats["总浏览数"]):,}')
        doc.add_paragraph(f'- 总点赞数：{int(stats["总点赞数"]):,}')
        doc.add_paragraph(f'- 总收藏数：{int(stats["总收藏数"]):,}')

        # 代表作列表
        doc.add_paragraph('代表作品：')
        for i, work in enumerate(author_top_works[author], 1):
            doc.add_paragraph(f'{i}. {work}', style='List Number')

        # 添加分隔线
        doc.add_paragraph('=' * 50)

    # 保存详细报告
    doc.save(output_dir / '作者分析详细报告.docx')

    # 将代表作信息添加到作者排名CSV中
    author_ranking['代表作品'] = author_ranking.index.map(
        lambda x: ' || '.join(author_top_works.get(x, ['无']))
    )

    # 更新输出列
    output_columns = ['排名', '文章数量', '平均文章评分', '总字数',
                     '总浏览数', '总点赞数', '总收藏数',
                     '质量稳定性', '作者评分', '代表作品']

    # 保存作者排名
    author_ranking[output_columns].to_csv(
        output_dir / '作者推荐排名.csv',
        encoding='utf-8-sig'
    )

    logging.info("作者分析完成，报告已生成")
    # 4. 标准化函数